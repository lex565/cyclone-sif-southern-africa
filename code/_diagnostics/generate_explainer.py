import os
# -*- coding: utf-8 -*-
"""
generate_explainer.py — detailed plain-English Word guide for each event folder.
Written as if explaining the results to a person: what each figure is, what every
axis / colour-bar number means (with units), what you actually SEE in this event
(data-driven), the interesting insights, and WHY the three-signal (SIF/NIRvR/PhiF)
approach is used. Replaces "de-meaned" with "normal-year adjusted".

Usage:
  python generate_explainer.py                 # all completed event folders
  python generate_explainer.py LABEL [LABEL..]  # specific event folders
Output: <event>/<LABEL>_RESULTS_EXPLAINED.docx
"""
import os, csv, glob, sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

RESULTS_ROOT = ros.environ.get("CYCLONE_SIF_OUT", r"derived_data")
WIN_ORDER = ["baseline","pre7","acute","early","recov2","recov3"]
WIN_PLAIN = {
    "baseline":"the two weeks BEFORE the storm (our reference)",
    "pre7":"the final week before landfall",
    "acute":"the storm week itself (landfall + 6 days)",
    "early":"the first week after the storm",
    "recov2":"two-to-three weeks after",
    "recov3":"three-to-four weeks after",
}

def parse_label(evdir):
    label = os.path.basename(evdir.rstrip("\\/")); parent = os.path.basename(os.path.dirname(evdir))
    known_dirs = ["Zimbabwe","Malawi","Botswana","Mozambique","Madagascar","South_Africa"]
    parts = label.split("_")
    year = parts[-1] if parts[-1].isdigit() else "?"
    if parent in known_dirs:                       # NESTED layout: Country/Storm_Year
        country = parent.replace("_"," "); storm = parts[0].title()
    else:                                          # FLAT label: PREFIX_STORM_YEAR
        known = {"MOZAMBIQUE":"Mozambique","MADAGASCAR":"Madagascar","SOUTH":"South Africa",
                 "BOTSWANA":"Botswana","MALAWI":"Malawi"}
        if parts[0] in known:
            country = known[parts[0]]; storm = parts[2].title() if parts[0]=="SOUTH" else parts[1].title()
        else:
            country = "Zimbabwe"; storm = parts[0].title()
    return country, storm, year

def read_rows(path):
    if not os.path.exists(path): return []
    with open(path, newline="", encoding="utf-8-sig") as f:
        return list(csv.DictReader(f))

def fnum(x, nd=2, d=None):
    try: return round(float(x), nd)
    except: return d

# ---- docx helpers ----
def para(doc, text, bold=False, italic=False, size=11, color=None, indent=0.0):
    p = doc.add_paragraph();
    if indent: p.paragraph_format.left_indent = Inches(indent)
    r = p.add_run(text); r.bold=bold; r.italic=italic; r.font.size=Pt(size)
    if color: r.font.color.rgb = RGBColor(*color)
    return p

def insight(doc, body):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.25)
    r = p.add_run("💡 Interesting insight: "); r.bold=True; r.font.size=Pt(10.5)
    r2 = p.add_run(body); r2.italic=True; r2.font.size=Pt(10.5)

def axisbox(doc, body):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.25)
    r = p.add_run("What the numbers mean — "); r.bold=True; r.font.size=Pt(10.5)
    r2 = p.add_run(body); r2.font.size=Pt(10.5)

def callout(doc, title, body):
    p = doc.add_paragraph(); r = p.add_run("▶ " + title); r.bold=True; r.font.size=Pt(11)
    p2 = doc.add_paragraph(); p2.paragraph_format.left_indent = Inches(0.3)
    r2 = p2.add_run(body); r2.italic=True; r2.font.size=Pt(10.5)

def fig(doc, path, caption):
    if not path or not os.path.exists(path): return False
    doc.add_heading(caption, level=2)
    try:
        doc.add_picture(path, width=Inches(6.4)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        para(doc, f"[could not embed image: {e}]", italic=True)
    return True

def gfirst(folder, pat):
    h = sorted(glob.glob(os.path.join(folder, pat))); return h[0] if h else None

# ---- describe a panel from the data ----
def dir_words(v, near=0.005):
    if v is None: return "had no usable data"
    if abs(v) < near: return "barely moved from normal — pale/near-white"
    return ("dropped below normal — reddish" if v < 0 else "rose above normal — greenish")

def build(evdir):
    label = os.path.basename(evdir.rstrip("\\/"))
    country, storm, year = parse_label(evdir)
    sm = os.path.join(evdir, "spatial_maps"); mt = os.path.join(evdir, "metrics"); at = os.path.join(evdir, "attribution")

    spat = {}
    for r in read_rows(os.path.join(sm, "spatial_measurements.csv")):
        spat[(r["window"], r["signal"])] = r
    win = {r["window"]: r for r in read_rows(os.path.join(evdir, "window_anomalies.csv"))}
    cov = read_rows(os.path.join(mt, "eq1_eq2_coverage_dilution.csv"))
    eq6 = {r["window"]: r for r in read_rows(os.path.join(mt, "eq6_pct_change.csv"))}
    eq7 = read_rows(os.path.join(mt, "eq7_recovery_ratios.csv"))
    eq8 = read_rows(os.path.join(mt, "eq8_sif_gpp_correlation.csv"))
    rcl = read_rows(os.path.join(mt, "response_class_summary.csv"))
    shn = read_rows(os.path.join(mt, "shannon_entropy.csv"))
    acc = spat.get(("acute","sif")); acn = spat.get(("acute","nirvr"))
    blackout = (acc is None) or (int(acc.get("n_cells","0") or 0) == 0)

    doc = Document()
    doc.add_heading(f"Cyclone {storm} over {country} ({year})", level=0)
    para(doc, "A detailed reading guide — written to be understood without a remote-sensing background.",
         italic=True, size=12)
    para(doc, f"Dear reader,\n\nThis note walks you through every figure for Cyclone {storm} in {country}: "
              f"what it shows, what the numbers on the axes and colour bars actually mean, what you should "
              f"see when you look at it, and why it matters. I keep the real science, but say it plainly. "
              f"Take it figure by figure — there's no rush.", size=11)

    # WHY THREE SIGNALS
    doc.add_heading("First: why we look at THREE signals, not one", level=1)
    para(doc, "Our main measurement is SIF — solar-induced fluorescence — a faint glow living plants give "
              "off while photosynthesising. Think of it as the plant's heartbeat. But there is a trap.")
    para(doc, "In a big region like this one, the land is a patchwork: dense woodland in some places, open "
              "savanna in others, cropland elsewhere. Raw SIF is naturally HIGH where there is simply a lot "
              "of leafy canopy, and LOW where vegetation is sparse — whether or not anything is stressed. So "
              "a raw-SIF map mixes up two different questions: 'how much vegetation is here?' and 'how hard "
              "is that vegetation working?' Over a mixed landscape, that mixing is misleading.")
    para(doc, "That is exactly why we add a second signal, NIRvR. NIRvR measures the 'how much green, lit "
              "canopy is present' part. When we divide SIF by NIRvR we get a third signal, ΦF (\"phi-F\"): "
              "the photosynthesis glow PER UNIT of canopy — the heartbeat per unit of body. ΦF is finally "
              "comparable between a thick forest pixel and a thin cropland pixel, because the amount of "
              "vegetation has been divided out.", )
    para(doc, "So the three signals answer three different questions:", bold=True)
    para(doc, "• SIF — total photosynthesis glow (structure × efficiency × light all mixed).", indent=0.25)
    para(doc, "• NIRvR — how much green canopy is physically present and lit (the 'body').", indent=0.25)
    para(doc, "• ΦF = SIF ÷ NIRvR — photosynthetic efficiency per unit canopy (the 'heartbeat per body').", indent=0.25)
    insight(doc, "This split lets us tell apart two very different kinds of damage. If NIRvR falls, the canopy "
                 "was PHYSICALLY destroyed — leaves stripped, vegetation flattened. If instead ΦF falls, the "
                 "surviving leaves are alive but photosynthesising less efficiently. A cyclone can do either, "
                 "and a single SIF number cannot tell them apart. That is the core reason for the three-signal "
                 "design — and a side-benefit is that it also handles the mixed-vegetation problem above.")

    # NORMAL-YEAR ADJUSTMENT
    doc.add_heading("Second: the 'normal-year adjustment' you'll see mentioned", level=1)
    off = fnum((win.get("baseline") or {}).get("sif_anom"), 3)
    para(doc, "Some years are naturally greener or drier than others before any storm arrives. If we ignored "
              "that, we might blame the cyclone for a head start it did not cause. So we measure how far this "
              "year already sat from a normal year in the days BEFORE the storm, and subtract that head start. "
              "Whatever change remains is genuinely the storm's doing.")
    if off is not None:
        para(doc, f"For this event the pre-storm head start was about {off:+.3f} SIF units; once removed, the "
                  f"'before' picture reads as zero and every later change is measured against a fair, normal "
                  f"baseline.", size=10.5)
    para(doc, "Analogy: to measure how much a storm lowered a river, first note the river was already low from "
              "a dry spell, subtract that, and the rest is the storm. Anywhere you read 'normal-year adjusted' "
              "(the technical files call it 'de-meaned'), this is all it means.", italic=True, size=10.5)

    # NO-DATA
    doc.add_heading("Third: the blank / grey areas (no-data, or 'blackout')", level=1)
    para(doc, "TROPOMI sees plants by their sunlight glow, and clouds block that view. A cyclone is a giant "
              "cloud system, so for several days during and after landfall the satellite literally cannot see "
              "the ground beneath the storm. Those days and places come back empty, and we leave them blank "
              "rather than guess.")
    para(doc, "A blank panel means 'the cloud hid it from us', NOT 'there was no damage'. This is a known limit "
              "of every optical satellite for cyclones — and it is itself a finding: the most violent moment of "
              "a storm is often the hardest to see from space. When the cloud clears (usually days later), the "
              "data returns.", size=10.5)
    if blackout:
        para(doc, f"For Cyclone {storm} over {country}, this blackout dominates the storm week — which is why "
                  f"several panels here are blank. We report it honestly. The maps below show mostly the "
                  f"before-and-after; the during is largely unseen.", bold=True, size=10.5)

    # FIGURE WALKTHROUGH
    doc.add_heading("Now the figures, one at a time", level=1)

    cfd = os.path.join(evdir, "clean_figures")   # the clean, supervisor-ready figures

    # --- Figure A: SIF over time (CLEAN) ---
    if fig(doc, gfirst(cfd, "*CLEAN_A_SIF_time.png"),
           "Figure A — SIF (photosynthesis) before, during, and after the storm"):
        axisbox(doc, "Three maps, left to right: the week BEFORE the storm, the STORM WEEK, and 3–4 WEEKS AFTER. "
                     "Colour = change versus a normal year, in seven bands from strong suppression (dark red, more "
                     "than 40% below normal) through no change (white, ±5%) to strong enhancement (dark green, more "
                     "than 40% above). The black line is the cyclone track; blank areas are cloud blackout. The maps "
                     "are smoothed for display; all statistics are computed on the original pixels.")
        if not blackout and eq6.get("acute"):
            ap=fnum(eq6["acute"]["pct_corrected"],0)
            if ap is not None:
                para(doc, f"What you see: in the storm-week panel, photosynthesis averaged about {abs(ap):.0f}% "
                          f"{'below' if ap<0 else 'above'} normal across the corridor; by 3–4 weeks after, green "
                          f"returns as the season greens up.", size=10.5)

    # --- Figure B: three signals, storm week (CLEAN) ---
    if fig(doc, gfirst(cfd, "*CLEAN_B_triplet_acute.png"),
           "Figure B — The three signals in the storm week (structure vs. physiology)"):
        axisbox(doc, "The same area mapped three times for the storm week: SIF (photosynthesis glow), NIRvR (how "
                     "much green canopy is present), and ΦF (efficiency = SIF ÷ NIRvR). Same seven-band colour scale "
                     "(% vs normal). This is the figure that separates physical canopy damage from a physiological "
                     "slowdown — the heart of the method.")
        if not blackout and acc and acn:
            sm_v=fnum(acc["demeaned_mean"],3); nv=fnum(acn["demeaned_mean"],2); sup=fnum(acn["pct_suppressed"],0)
            para(doc, f"What you see: SIF (left) {dir_words(sm_v)}, but NIRvR (middle) {dir_words(nv,0.3)} with about "
                      f"{sup:.0f}% of locations below normal — the canopy structure took the hit.", size=10.5)
            if sm_v is not None and nv is not None and abs(sm_v) < 0.05 and nv < -0.5:
                insight(doc, "The striking part: SIF barely moved while NIRvR fell sharply. SIF alone would have "
                             "said 'no impact'. NIRvR reveals the canopy was physically stripped while the surviving "
                             "leaves kept their efficiency — STRUCTURAL damage, not a physiological shutdown. This is "
                             "the single best argument for the three-signal method.")

    # --- Response classes ---
    if fig(doc, gfirst(mt, "*response_classes_acute.png"), "Figure — Storm-week response map (five categories)"):
        axisbox(doc, "Each location is sorted into one of five colours, from strong suppression (deep red, more "
                     "than 20% below normal) through negligible (pale) to strong enhancement (deep green, more "
                     "than 20% above). There is no number axis here — it is a category map.")
        if rcl:
            ss = next((r for r in rcl if "Strong suppression" in r["class"]), None)
            se = next((r for r in rcl if "Strong enhancement" in r["class"]), None)
            ev = fnum((shn[0]["evenness"] if shn else None),2)
            txt = []
            if ss: txt.append(f"{float(ss['proportion'])*100:.0f}% strongly suppressed")
            if se: txt.append(f"{float(se['proportion'])*100:.0f}% strongly enhanced")
            para(doc, "What you see: " + ", ".join(txt) + f", and a mixedness score of {ev}.", size=10.5)
            callout(doc, "What the 'mixedness score' (Shannon evenness) means, in plain English",
                    f"Shannon evenness is just a number from 0 to 1 that says how VARIED the responses were across "
                    f"the area. Think of five buckets (strong suppression → strong enhancement) and imagine "
                    f"dropping every pixel into its bucket. If they all land in ONE bucket, the score is 0 — the "
                    f"whole area reacted the same way (uniform). If they spread EVENLY across all five buckets, the "
                    f"score is 1 — maximally mixed/patchy. Here the score is {ev}, "
                    f"{'so the storm’s effect was very uneven — it hit some places hard while others nearby were barely touched or even greener.' if ev and ev>0.8 else 'so the response was fairly uniform across the area.'} "
                    f"(The technical name is Shannon entropy; evenness is that entropy divided by its maximum so it "
                    f"always sits between 0 and 1.)")
            insight(doc, "A single corridor-average would hide this completely. Strong suppression and strong "
                         "enhancement can sit side by side — which is why we map every pixel and measure the "
                         "spread, instead of reporting one number for the whole country.")

    # --- Attribution ---
    if fig(doc, gfirst(at, "*driver_maps.png"), "Figure — What drove the damage (cause maps)"):
        axisbox(doc, "Four small maps. Far left: the measured storm-week SIF change (red = loss). The other "
                     "three are the suspected causes — closeness to the storm track (a wind-exposure proxy), how "
                     "wet the soil already was (antecedent 60-day rain, in mm), and the storm's own rainfall "
                     "(mm). You compare the cause maps to the damage map to see which lines up.")
    cf_dir = os.path.join(evdir, "clean_figures")
    if fig(doc, gfirst(cf_dir, "*CLEAN_attribution.png"), "Figure — The same causes, as scatter plots"):
        axisbox(doc, "Each dot is one location. The horizontal axis is a cause — antecedent (60-day) rainfall in "
                     "mm, the storm's own rainfall in mm, or wind exposure (closeness to the track). The vertical "
                     "axis is the storm-week SIF change as % of normal. A clear downhill or uphill tilt (and the r "
                     "value above each panel) means that cause is linked to the damage; a shapeless blob means it "
                     "is not. r runs from -1 to +1; near 0 means no link.")
        insight(doc, "The two rainfall panels test whether wetter ground (before) or heavier storm rain "
                     "lines up with the vegetation change; the wind panel tests proximity to the track. Read the "
                     "tilt, not individual dots — the regression line is the summary.")
        coefs = read_rows(os.path.join(at, "attribution_coefficients.csv"))
        sig = [c for c in coefs if c.get("sig","")=="*"]
        if sig:
            names = {"wind_stress":"closeness to the track (wind)","ante_moist":"how wet the soil already was",
                     "acute_rain":"the storm's rainfall","wind_stress:ante_moist":"wind AND wet soil together"}
            terms = ", ".join(names.get(c["term"], c["term"]) for c in sig if c["term"]!="(Intercept)")
            if terms:
                insight(doc, f"For this event the factor(s) that significantly matched the damage pattern: {terms}. "
                             f"Where you see 'wind AND wet soil together', it means the worst damage happened "
                             f"where strong winds hit ground that was already saturated — the two compounding.")

    # --- Coverage ---
    if fig(doc, gfirst(sm, "*coverage_per_window.png"), "Figure — How much we could actually see"):
        axisbox(doc, "Bars (or counts) per time window showing how many valid satellite readings we had. The "
                     "horizontal axis is the time window; the vertical axis is the number of clear readings. "
                     "Low or empty bars in the storm week are the cloud blackout described earlier — they tell "
                     "you how much to trust each column in the maps above.")

    # NUMBERS IN PLAIN WORDS
    doc.add_heading("The headline numbers, in plain words", level=1)
    if cov:
        cf=fnum(cov[0]["CF"],2); df=fnum(cov[0]["DF"],1)
        para(doc, f"• Footprint coverage (Eq 1 & 2): only ~{cf*100:.0f}% of {country} lay inside the storm's "
                  f"200 km footprint. Averaging the whole country would dilute the signal ~{df:.0f}× — which is "
                  f"why we zoom into the corridor the storm actually crossed.")
    if not blackout and eq6.get("acute"):
        ap=fnum(eq6["acute"]["pct_corrected"],0)
        if ap is not None:
            para(doc, f"• Storm-week photosynthesis (Eq 6, corridor average): about {abs(ap):.0f}% "
                      f"{'below' if ap<0 else 'above'} normal for the season.")
    if eq7:
        late = next((r for r in eq7 if r.get("phase")=="late_recovery"), None)
        rv = fnum(late.get("RR_vs_clim"),2) if late and late.get("RR_vs_clim") not in (None,"","NA") else None
        if rv is not None:
            para(doc, f"• Recovery 6–8 weeks later (Eq 7): ~{rv*100:.0f}% of the normal seasonal level "
                      f"({'essentially recovered' if rv>=0.95 else 'still suppressed'}).")
    if eq8:
        r=fnum(eq8[0]["pearson_r"],2); n=eq8[0].get("n_pairs","?")
        para(doc, f"• Independent GPP check (Eq 8): agreement r={r} over n={n} periods"
                  f"{' (small n — indicative only)' if str(n).isdigit() and int(n)<5 else ''}.")
    # else: GPP not available for this country — skip the GPP validation entirely (omitted).

    # WHICH EQUATION DID WHAT
    doc.add_heading("Which equation calculated what", level=1)
    para(doc, "Every number and map above comes from a specific equation in the Methods. Here is the "
              "plain-language map between them, so you can point a reader (or examiner) straight to the source:")
    table = doc.add_table(rows=1, cols=3); table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Equation"; hdr[1].text = "What it produces here"; hdr[2].text = "Where you see it"
    eqrows = [
        ("Eq 1 — Coverage fraction", "What share of the country sat inside the storm footprint", "Footprint coverage % (headline numbers)"),
        ("Eq 2 — Dilution factor", "How much a nationwide average would water down the signal", "The '~N× dilution' number"),
        ("Eq 5 — Anomaly", "Observed SIF minus the normal-year value, per grid cell", "Every map (the colours ARE the anomaly)"),
        ("Eq 6 — % change vs baseline", "Storm-week change as a % of the normal level", "Storm-week photosynthesis % ; the map colour classes"),
        ("Eq 7 — Recovery ratio", "How far function returned weeks later, vs the seasonal normal", "The 6–8-week recovery number"),
        ("Eq 8 — Pearson r (SIF vs GPP)", "Whether SIF agrees with an independent productivity product", "Independent GPP check (deferred where GPP not yet downloaded)"),
        ("Eq 9 — Attribution regression", "Which drivers (wind / antecedent rain / storm rain) match the damage", "The attribution cause maps and scatter plots"),
        ("§3.7 — 5-class + Shannon", "Sorting each pixel into suppression/enhancement classes + a mixedness score", "The storm-week response map and the variety-of-responses number"),
        ("NIRvR / ΦF (Zeng 2022)", "Splitting structure (NIRvR) from efficiency (ΦF = SIF/NIRvR)", "The middle and bottom maps of the triplet figure"),
    ]
    for a,b,c in eqrows:
        cells = table.add_row().cells; cells[0].text=a; cells[1].text=b; cells[2].text=c
        for cell in cells:
            for p in cell.paragraphs:
                for r in p.runs: r.font.size = Pt(9.5)

    doc.add_heading("Honest limitations", level=1)
    para(doc, "• Cloud blackout limits the view during the storm itself.\n"
              "• The satellite footprint is a few km, so very local damage can be averaged out.\n"
              "• ΦF is noisy when the signal is weak, so it is used only as a supporting clue.\n"
              "• Weeks-later 'recovery' is partly normal seasonal greening, which we remove by comparing to "
              "the same dates in normal years.", size=10.5)
    para(doc, "\n— End of guide. If any single figure still feels unclear, tell me which one and I will expand it.",
         italic=True, size=10.5)

    out = os.path.join(evdir, f"{label}_RESULTS_EXPLAINED.docx")
    try:
        doc.save(out)
    except PermissionError:
        out = os.path.join(evdir, f"{label}_RESULTS_EXPLAINED_NEW.docx")
        doc.save(out)  # original is open in Word; write a fresh copy instead
    return out, ("BLACKOUT" if blackout else "OK")

def main():
    args = sys.argv[1:]
    if args:
        dirs = [os.path.join(RESULTS_ROOT, a) for a in args]
    else:
        # recurse: an event folder has window_anomalies.csv + <name>_original_metrics.md (works flat OR nested)
        dirs = []
        for root, subs, files in os.walk(RESULTS_ROOT):
            if "window_anomalies.csv" in files and any(f.endswith("_original_metrics.md") for f in files):
                dirs.append(root)
    print(f"Generating detailed explainers for {len(dirs)} event(s)...")
    for d in sorted(dirs):
        try:
            out, st = build(d); print(f"  [{st}] {os.path.basename(out)}")
        except Exception as e:
            import traceback; print(f"  !! FAILED {os.path.basename(d)}: {e}"); traceback.print_exc()
    print("DONE.")

if __name__ == "__main__":
    main()
